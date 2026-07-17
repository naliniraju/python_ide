"""
Python IDE - Streamlit Mini-IDE
--------------------------------------
Browse a local folder, pick a .py or .ipynb file, edit it, run it, send
stdin input to it while it runs, and see live output + errors.

Run with:
    streamlit run python_ide.py
"""

import os
import sys
import base64
import shutil
import select
import logging
import platform
import subprocess
from pathlib import Path
from dataclasses import dataclass
from logging.handlers import RotatingFileHandler

import streamlit as st
import streamlit.components.v1 as components

try:
    from streamlit_ace import st_ace
    ACE_AVAILABLE = True
except ImportError:
    ACE_AVAILABLE = False

try:
    import nbformat
    from nbclient import NotebookClient
    from nbclient.exceptions import CellExecutionError
    NBCLIENT_AVAILABLE = True
except ImportError:
    NBCLIENT_AVAILABLE = False


# ============================= App configuration ============================
@dataclass(frozen=True)
class AppConfig:
    """Central place for the constants that shape app behaviour, instead of
    magic numbers scattered through the file — the standard "config object"
    pattern for production code."""
    APP_NAME: str = "Python IDE"
    APP_ICON: str = "🧪"
    VERSION: str = "2.0.0"
    MAX_OUTPUT_CHARS: int = 100_000       # console output cap, keeps re-renders fast
    CONSOLE_POLL_SECONDS: float = 0.5     # live-output refresh cadence while a process runs
    SUBPROCESS_LAUNCH_TIMEOUT: int = 180  # native folder-picker subprocess timeout (seconds)
    LOG_DIR: Path = Path.home() / ".python_ide" / "logs"
    LOG_FILE: str = "app.log"
    LOG_MAX_BYTES: int = 2_000_000
    LOG_BACKUP_COUNT: int = 3


CONFIG = AppConfig()


# ============================= Logging =======================================
def _setup_logger() -> logging.Logger:
    """Configure a rotating file logger. Streamlit reruns the whole script
    on every interaction, so handlers are only attached once (guarded by
    `logger.handlers`) to avoid duplicate log lines piling up rerun after
    rerun."""
    logger = logging.getLogger("python_ide")
    logger.setLevel(logging.INFO)
    if not logger.handlers:
        try:
            CONFIG.LOG_DIR.mkdir(parents=True, exist_ok=True)
            handler = RotatingFileHandler(
                CONFIG.LOG_DIR / CONFIG.LOG_FILE,
                maxBytes=CONFIG.LOG_MAX_BYTES,
                backupCount=CONFIG.LOG_BACKUP_COUNT,
            )
            handler.setFormatter(logging.Formatter(
                "%(asctime)s [%(levelname)s] %(message)s", datefmt="%Y-%m-%d %H:%M:%S",
            ))
            logger.addHandler(handler)
        except OSError:
            # Fall back to console-only logging if the log dir can't be created
            # (e.g. read-only home directory) — logging should never crash the app.
            logger.addHandler(logging.StreamHandler())
    return logger

def _get_python_executable() -> str:
    """Return the Python interpreter to use for running user scripts.

    When frozen by PyInstaller, sys.executable points at the frozen app
    binary itself (not a real Python interpreter) - so subprocess calls
    using it directly would just try to relaunch this app with the
    script's path as a garbage argument, which fails immediately.
    Fall back to whatever python3 is on PATH in that case.
    """
    if getattr(sys, "frozen", False):
        return shutil.which("python3") or shutil.which("python") or "python3"
    return sys.executable
logger = _setup_logger()

if NBCLIENT_AVAILABLE:
    # Under PyInstaller, jupyter_client's entry-point-based plugin lookup
    # for kernel provisioners ("create_provisioner_instance") breaks with
    # "TypeError: 'module' object is not callable" - a known rough edge
    # between importlib.metadata entry-points and frozen apps. This app
    # only ever needs the built-in local provisioner (a plain local
    # Python kernel), so we bypass the broken plugin lookup entirely and
    # force it to always use LocalProvisioner directly.
    try:
        from jupyter_client.provisioning.local_provisioner import LocalProvisioner
        from jupyter_client.provisioning.factory import KernelProvisionerFactory

        def _patched_create_provisioner_instance(self, kernel_id, kernel_spec, parent=None):
            provisioner_config = {}
            if kernel_spec.metadata and "kernel_provisioner" in kernel_spec.metadata:
                provisioner_config = kernel_spec.metadata["kernel_provisioner"].get("config", {})
            return LocalProvisioner(
                kernel_id=kernel_id, kernel_spec=kernel_spec, parent=parent, **provisioner_config
            )

        KernelProvisionerFactory.create_provisioner_instance = _patched_create_provisioner_instance
    except Exception:
        logger.exception("Could not patch jupyter_client provisioner factory")

st.set_page_config(page_title=CONFIG.APP_NAME, page_icon=CONFIG.APP_ICON, layout="wide")
# ============================= Constants ====================================
SKIP_DIRS = {
    ".git", "__pycache__", ".ipynb_checkpoints", "node_modules",
    ".venv", "venv", ".mypy_cache", ".pytest_cache",
}

RUNNABLE_EXTS = {".py", ".ipynb"}

# Extensions that can be clicked in the tree to open a read-only preview
# (no editing/running - just a quick look at the file's content).
PREVIEWABLE_EXTS = {
    ".pdf", ".docx", ".txt", ".xlsx", ".csv", ".png", ".jpg", ".jpeg", ".svg",
}

# Only two icons are used, deliberately — one per runnable type. A
# per-extension emoji map for every other file type (.xlsx, .json, .yml...)
# looks appealing in theory but real-world font stacks render many of
# those glyphs inconsistently (missing glyph boxes, mismatched sizes,
# oddly colored fallbacks), which is what made the tree look broken.
# Non-runnable files intentionally get no icon at all — just a small
# muted bullet — so the tree stays clean regardless of the system font.
RUNNABLE_ICONS = {
    ".py": "🐍",
    ".ipynb": "📓",
}

# Small, deliberately coarse icon set for previewable files — grouped by
# broad kind (document / tabular / image) rather than one glyph per
# extension, for the same font-consistency reason as above.
PREVIEW_ICONS = {
    ".pdf": "📄",
    ".docx": "📄",
    ".txt": "📄",
    ".xlsx": "📊",
    ".csv": "📊",
    ".png": "🖼️",
    ".jpg": "🖼️",
    ".jpeg": "🖼️",
    ".svg": "🖼️",
}


def _default_root_folder() -> str:
    """Default the file browser to the user's Desktop if it exists,
    otherwise fall back to their home directory."""
    desktop = Path.home() / "Desktop/py_source_83"
    if desktop.is_dir():
        return str(desktop)
    return str(Path.home())


# ============================= Session state ==============================
defaults = {
    "folder_path": _default_root_folder(),
    "selected_file": None,
    "editor_content": {},      # filepath -> current (possibly edited) text
    "original_content": {},    # filepath -> content as last loaded from disk
    "proc": None,              # running subprocess.Popen, or None
    "proc_output": "",
    "proc_returncode": None,
    "proc_file": None,
    "nb_sessions": {},         # filepath -> {"nb": NotebookNode, "client": NotebookClient, "kernel_started": bool}
    "nb_cell_source": {},      # (filepath, cell_index) -> current edited source text
    "nb_cell_original": {},    # (filepath, cell_index) -> source as last saved to disk
    "fs_version": 0,           # bumped whenever the filesystem changes; invalidates cached folder scans
    "editor_version": {},      # filepath -> int, bumped to force the code editor to remount
    "cell_editor_version": {}, # (filepath, cell_index) -> int, same idea per notebook cell
    "expanded_folders": set(), # rel folder paths ("" = root) currently expanded in the tree
    "show_inline_new_file": False,
    "show_inline_new_folder": False,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v


def ace_height_for(text: str, min_lines: int = 5, line_height: int = 19, padding: int = 34) -> int:
    """Compute an editor height tall enough to show all lines of `text` with
    no internal scrollbar, instead of using a fixed height."""
    lines = text.count("\n") + 1
    lines = max(lines, min_lines)
    return lines * line_height + padding


class PathError(ValueError):
    """Raised when a requested file/folder path is invalid or unsafe."""


def _resolve_safe_path(root_folder: str, rel_path: str) -> Path:
    """Resolve `rel_path` against `root_folder` and guarantee the result
    cannot escape the root (blocks '..', symlink tricks, absolute paths
    pointing elsewhere, etc). Raises PathError with a user-facing message
    on any problem.

    An absolute path is accepted as a convenience (e.g. pasted from a
    file manager or from the sidebar's own path display) as long as it
    already resolves to somewhere inside root_folder; only absolute paths
    that point outside the root are rejected.
    """
    rel_path = (rel_path or "").strip()
    if not rel_path:
        raise PathError("Please enter a name or path.")

    root = Path(root_folder).expanduser().resolve()

    if os.path.isabs(rel_path):
        candidate = Path(rel_path).expanduser().resolve()
    else:
        candidate = (root / rel_path).resolve()

    try:
        candidate.relative_to(root)
    except ValueError:
        raise PathError(
            "That path is outside the root folder "
            f"({root}) — use a path relative to it, or an absolute path inside it."
        )

    if candidate == root:
        raise PathError("You can't target the root folder itself.")

    return candidate


def create_new_file(root_folder: str, rel_path: str) -> tuple[bool, str]:
    """Create a new file at rel_path (relative to root_folder).

    Notebook files (.ipynb) are initialized with a valid, empty notebook
    document rather than an empty string, since an empty file isn't valid
    JSON and would fail to open (`nbformat.read` would raise).
    Returns (success, absolute_path_or_error_message).
    """
    try:
        target = _resolve_safe_path(root_folder, rel_path)
    except PathError as e:
        return False, str(e)

    if target.exists():
        return False, "A file or folder already exists at that path."

    try:
        target.parent.mkdir(parents=True, exist_ok=True)
        if target.suffix == ".ipynb" and NBCLIENT_AVAILABLE:
            nbformat.write(nbformat.v4.new_notebook(), str(target))
        else:
            target.write_text("", encoding="utf-8")
        logger.info("Created file: %s", target)
        return True, str(target)
    except OSError as e:
        logger.exception("Failed to create file: %s", target)
        return False, f"Couldn't create file: {e}"


def create_new_folder(root_folder: str, rel_path: str) -> tuple[bool, str]:
    """Create a new (possibly nested) folder at rel_path (relative to root_folder)."""
    try:
        target = _resolve_safe_path(root_folder, rel_path)
    except PathError as e:
        return False, str(e)

    if target.exists():
        return False, "A file or folder already exists at that path."

    try:
        target.mkdir(parents=True, exist_ok=False)
        logger.info("Created folder: %s", target)
        return True, str(target)
    except OSError as e:
        logger.exception("Failed to create folder: %s", target)
        return False, f"Couldn't create folder: {e}"


def delete_path(root_folder: str, rel_path: str) -> tuple[bool, str]:
    """Delete a file or folder at rel_path (relative to root_folder)."""
    try:
        target = _resolve_safe_path(root_folder, rel_path)
    except PathError as e:
        return False, str(e)

    if not target.exists():
        return False, "That path doesn't exist."

    try:
        if target.is_dir():
            shutil.rmtree(target)
        else:
            target.unlink()
        logger.info("Deleted: %s", target)
        return True, str(target)
    except OSError as e:
        logger.exception("Failed to delete: %s", target)
        return False, f"Couldn't delete: {e}"


def _browse_for_folder_macos(initial_dir: str) -> tuple[bool, str]:
    """Native macOS folder picker via AppleScript's `choose folder`, run
    through `osascript` in its own process.

    This deliberately does NOT use tkinter in-process: Tkinter's dialogs on
    macOS call into Cocoa (NSWindow/NSOpenPanel), and Cocoa APIs may only
    be touched from the main thread of a process. Streamlit runs the
    script inside a background ScriptRunner thread, not the main thread,
    so an in-process Tkinter dialog crashes the whole app with
    `libc++abi ... NSException`. Shelling out to `osascript` sidesteps
    this entirely — it's a separate OS process with its own main thread,
    so Cocoa is safe to call there regardless of which thread launched it.
    """
    start_dir = initial_dir if os.path.isdir(initial_dir) else str(Path.home())
    script = (
        f'set startFolder to POSIX file "{start_dir}"\n'
        'set chosenFolder to choose folder with prompt "Select a folder" default location startFolder\n'
        'return POSIX path of chosenFolder'
    )
    try:
        result = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=CONFIG.SUBPROCESS_LAUNCH_TIMEOUT,
        )
    except Exception as e:
        logger.exception("macOS folder picker failed")
        return False, f"Couldn't open the folder picker: {e}"

    if result.returncode != 0:
        stderr = (result.stderr or "").strip()
        if "User canceled" in stderr or "(-128)" in stderr:
            return False, ""  # user cancelled the dialog - not an error
        return False, f"Couldn't open the folder picker: {stderr or 'unknown error'}"

    chosen = result.stdout.strip()
    if not chosen:
        return False, ""
    return True, chosen


def _browse_for_folder_subprocess(initial_dir: str) -> tuple[bool, str]:
    """Fallback folder picker for non-macOS platforms: runs Tkinter's
    askdirectory() inside a brand-new Python subprocess (which gets its
    own main thread), rather than importing tkinter into this process.
    Requires a display to be available (won't work on a headless server)."""
    start_dir = initial_dir if os.path.isdir(initial_dir) else str(Path.home())
    code = (
        "import tkinter, tkinter.filedialog as fd, sys\n"
        "root = tkinter.Tk()\n"
        "root.withdraw()\n"
        "root.attributes('-topmost', True)\n"
        f"path = fd.askdirectory(initialdir={start_dir!r}, title='Select a folder')\n"
        "sys.stdout.write(path)\n"
    )
    try:
        result = subprocess.run(
            [sys.executable, "-c", code],
            capture_output=True, text=True, timeout=CONFIG.SUBPROCESS_LAUNCH_TIMEOUT,
        )
    except Exception as e:
        logger.exception("Subprocess folder picker failed")
        return False, f"Couldn't open the folder picker: {e}"

    if result.returncode != 0:
        stderr = (result.stderr or "").strip()
        return False, f"Couldn't open the folder picker: {stderr or 'unknown error'}"

    chosen = result.stdout.strip()
    if not chosen:
        return False, ""  # user cancelled the dialog - not an error
    return True, chosen


def browse_for_folder(initial_dir: str) -> tuple[bool, str]:
    """Open a native OS folder-picker dialog rooted at initial_dir, and
    return (True, chosen_path) or (False, error_or_empty_if_cancelled).

    This only works when the Streamlit process itself has access to a
    display (i.e. running locally on your own machine) — it will not
    work on a headless remote server.
    """
    if platform.system() == "Darwin":
        return _browse_for_folder_macos(initial_dir)
    return _browse_for_folder_subprocess(initial_dir)


# ============================= Helpers =====================================
@st.cache_data(show_spinner=False)
def _scan_runnable_files(folder: str, _fs_version: int):
    """Recursively list .py and .ipynb files under folder, skipping junk
    dirs. Cached via st.cache_data keyed on (folder, fs_version) so a
    repeated rerun (button clicks, cell runs, stdin sends, etc.) doesn't
    re-walk the filesystem — only a real change bumps fs_version and busts
    the cache. Returns (results, error_message_or_None)."""
    results = []
    try:
        for root, dirs, files in os.walk(folder):
            dirs[:] = [d for d in dirs if d not in SKIP_DIRS and not d.startswith(".")]
            for f in files:
                if f.endswith(".py") or f.endswith(".ipynb"):
                    full = os.path.join(root, f)
                    rel = os.path.relpath(full, folder)
                    results.append((rel, full))
    except Exception as e:
        logger.exception("Failed to scan runnable files under %s", folder)
        return [], f"Could not read folder: {e}"
    return sorted(results, key=lambda x: x[0].lower()), None


@st.cache_data(show_spinner=False)
def _scan_file_tree(folder: str, _fs_version: int):
    """Build a nested dict representing the full folder/file tree under
    `folder`, VS Code style (all files shown, not just runnable ones).
    Cached the same way as _scan_runnable_files.

    Shape: {"__folders__": {name: subtree, ...}, "__files__": [(name, full_path), ...]}
    Skips the same junk dirs as _scan_runnable_files and hides dotfiles.
    Returns (tree, error_message_or_None).
    """
    tree = {"__folders__": {}, "__files__": []}
    try:
        entries = sorted(os.listdir(folder), key=str.lower)
    except Exception as e:
        logger.exception("Failed to build file tree for %s", folder)
        return tree, f"Could not read folder: {e}"

    for name in entries:
        if name.startswith("."):
            continue
        full = os.path.join(folder, name)
        if os.path.isdir(full):
            if name in SKIP_DIRS:
                continue
            subtree, _ = _scan_file_tree(full, _fs_version)
            tree["__folders__"][name] = subtree
        else:
            tree["__files__"].append((name, full))
    return tree, None


def list_runnable_files(folder: str, fs_version: int):
    """Cache-aware wrapper around _scan_runnable_files that surfaces any
    error to the user via st.error (kept out of the cached function itself
    so the error still shows correctly even on a cache hit)."""
    results, error = _scan_runnable_files(folder, fs_version)
    if error:
        st.error(error)
    return results


def build_file_tree(folder: str, fs_version: int):
    """Cache-aware wrapper around _scan_file_tree — see list_runnable_files."""
    tree, error = _scan_file_tree(folder, fs_version)
    if error:
        st.error(error)
    return tree


def get_file_icon(name: str) -> str:
    ext = os.path.splitext(name)[1].lower()
    if ext in RUNNABLE_ICONS:
        return RUNNABLE_ICONS[ext]
    return PREVIEW_ICONS.get(ext, "")


def load_file_content(filepath: str) -> str:
    """Load (and cache) a file's text content."""
    if filepath not in st.session_state.original_content:
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                text = f.read()
        except Exception as e:
            text = f"# Could not read file: {e}"
        st.session_state.original_content[filepath] = text
        st.session_state.editor_content[filepath] = text
    return st.session_state.editor_content[filepath]


def save_file_content(filepath: str, content: str):
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)
    st.session_state.original_content[filepath] = content
    st.session_state.editor_content[filepath] = content


def kill_current_process():
    proc = st.session_state.proc
    if proc is not None and proc.poll() is None:
        try:
            proc.terminate()
            proc.wait(timeout=2)
            logger.info("Terminated process pid=%s for %s", proc.pid, st.session_state.proc_file)
        except Exception:
            try:
                proc.kill()
                logger.warning("Force-killed unresponsive process pid=%s for %s", proc.pid, st.session_state.proc_file)
            except Exception:
                logger.exception("Failed to kill process pid=%s", getattr(proc, "pid", "?"))
    st.session_state.proc = None
    st.session_state.proc_output = ""
    st.session_state.proc_returncode = None
    st.session_state.proc_file = None


def start_process(filepath: str):
    """Save current editor content to disk, then launch it as a subprocess."""
    kill_current_process()  # only one running process at a time

    cmd = [_get_python_executable(), "-u", filepath]
    workdir = os.path.dirname(filepath) or "."

    try:
        proc = subprocess.Popen(
            cmd,
            cwd=workdir,
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            bufsize=1,
        )
    except Exception as e:
        logger.exception("Failed to launch process for %s", filepath)
        st.session_state.proc_output = f"Failed to launch process: {e}"
        st.session_state.proc_returncode = -1
        return

    logger.info("Started process pid=%s for %s", proc.pid, filepath)
    st.session_state.proc = proc
    st.session_state.proc_output = ""
    st.session_state.proc_returncode = None
    st.session_state.proc_file = filepath


def _append_output(text: str):
    st.session_state.proc_output += text
    if len(st.session_state.proc_output) > CONFIG.MAX_OUTPUT_CHARS:
        trimmed = len(st.session_state.proc_output) - CONFIG.MAX_OUTPUT_CHARS
        st.session_state.proc_output = (
            f"...[{trimmed} earlier characters truncated]...\n"
            + st.session_state.proc_output[-CONFIG.MAX_OUTPUT_CHARS:]
        )


def poll_process_output():
    """Non-blocking read of any available output from the running process."""
    proc = st.session_state.proc
    if proc is None:
        return
    try:
        while True:
            readable, _, _ = select.select([proc.stdout], [], [], 0)
            if not readable:
                break
            chunk = os.read(proc.stdout.fileno(), 4096)
            if not chunk:
                break
            _append_output(chunk.decode(errors="replace"))
    except Exception:
        logger.exception("Error while polling process output for pid=%s", getattr(proc, "pid", "?"))

    if proc.poll() is not None and st.session_state.proc_returncode is None:
        try:
            remaining = proc.stdout.read()
            if remaining:
                _append_output(remaining)
        except Exception:
            logger.exception("Error reading remaining output for pid=%s", getattr(proc, "pid", "?"))
        st.session_state.proc_returncode = proc.returncode
        logger.info("Process pid=%s for %s exited with code %s", proc.pid, st.session_state.proc_file, proc.returncode)


def send_stdin(text: str):
    proc = st.session_state.proc
    if proc is not None and proc.poll() is None and proc.stdin:
        try:
            proc.stdin.write(text + "\n")
            proc.stdin.flush()
            _append_output(f"\n>>> {text}\n")
        except Exception as e:
            _append_output(f"\n[Could not send input: {e}]\n")


# ============================= Notebook execution ==========================
def render_notebook_output(output, container):
    otype = output.get("output_type")
    if otype == "stream":
        container.code(output.get("text", ""), language="text")
    elif otype in ("execute_result", "display_data"):
        data = output.get("data", {})
        if "image/png" in data:
            container.image(base64.b64decode(data["image/png"]))
        elif "image/jpeg" in data:
            container.image(base64.b64decode(data["image/jpeg"]))
        elif "text/html" in data:
            html = data["text/html"]
            container.markdown("".join(html) if isinstance(html, list) else html, unsafe_allow_html=True)
        elif "text/plain" in data:
            text = data["text/plain"]
            container.code("".join(text) if isinstance(text, list) else text, language="text")
    elif otype == "error":
        import re
        ename = output.get("ename", "Error")
        evalue = output.get("evalue", "")
        clean_tb = [re.sub(r"\x1b\[[0-9;]*m", "", line) for line in output.get("traceback", [])]
        container.error(f"{ename}: {evalue}")
        if clean_tb:
            container.code("\n".join(clean_tb), language="text")


def get_or_create_nb_session(filepath: str):
    """Load (or fetch cached) notebook + its NotebookClient. Kernel is NOT
    started here - it starts lazily on first cell run so opening a notebook
    is instant."""
    sessions = st.session_state.nb_sessions
    if filepath not in sessions:
        nb = nbformat.read(filepath, as_version=4)
        workdir = os.path.dirname(filepath) or "."
        client = NotebookClient(nb, timeout=600, kernel_name="python3", resources={"metadata": {"path": workdir}})
        sessions[filepath] = {"nb": nb, "client": client, "kernel_started": False}
        for idx, cell in enumerate(nb.cells):
            if cell.cell_type == "code":
                st.session_state.nb_cell_source.setdefault((filepath, idx), cell.source)
                st.session_state.nb_cell_original.setdefault((filepath, idx), cell.source)
    return sessions[filepath]


def ensure_kernel_started(session: dict):
    client = session["client"]
    if not session["kernel_started"]:
        client.km = client.create_kernel_manager()
        client.start_new_kernel()
        client.start_new_kernel_client()
        session["kernel_started"] = True


def restart_kernel(filepath: str):
    sessions = st.session_state.nb_sessions
    session = sessions.get(filepath)
    if session and session["kernel_started"]:
        try:
            session["client"]._cleanup_kernel()
        except Exception:
            pass
    if filepath in sessions:
        del sessions[filepath]


def run_single_cell(session: dict, cell_index: int, edited_source: str):
    """Apply the latest edited source to the cell, then execute it against
    the notebook's persistent kernel (starting the kernel on first use)."""
    nb = session["nb"]
    client = session["client"]
    cell = nb.cells[cell_index]
    cell.source = edited_source

    ensure_kernel_started(session)

    try:
        client.execute_cell(cell, cell_index)
        return True
    except CellExecutionError:
        return False
    except Exception as e:
        cell.outputs = [{
            "output_type": "error",
            "ename": "RunnerError",
            "evalue": str(e),
            "traceback": [str(e)],
        }]
        return False


def save_notebook(filepath: str, session: dict):
    nbformat.write(session["nb"], filepath)


def is_comment_only(source: str) -> bool:
    """True if every non-blank line in `source` is a comment (or the cell
    is empty) — i.e. there's no actual code to run or save."""
    for line in source.splitlines():
        stripped = line.strip()
        if stripped and not stripped.startswith("#"):
            return False
    return True


def _reset_cached_cell_state(filepath: str):
    """Drop all per-cell cached state for this notebook file so it gets
    freshly re-seeded from the notebook's actual (post-edit) content on
    the next render. Needed any time cell indices shift (insert/delete),
    since cached state is keyed by (filepath, cell_index)."""
    for state_dict in (
        st.session_state.nb_cell_source,
        st.session_state.nb_cell_original,
        st.session_state.cell_editor_version,
    ):
        for key in [k for k in state_dict if k[0] == filepath]:
            del state_dict[key]


def delete_notebook_cell(filepath: str, session: dict, cell_index: int):
    """Remove a cell from the notebook and persist the change to disk.

    Per-cell cached state (source buffer, original-on-disk copy, editor
    remount version) is keyed by (filepath, cell_index). Since removing a
    cell shifts every later cell's index down by one, those cached entries
    would otherwise point at the wrong cell — so we drop all cached state
    for this file and let it be freshly re-seeded from the notebook's
    actual (post-delete) content on the next render.
    """
    nb = session["nb"]
    if 0 <= cell_index < len(nb.cells):
        del nb.cells[cell_index]
    save_notebook(filepath, session)
    _reset_cached_cell_state(filepath)


def add_notebook_cell(filepath: str, session: dict, after_index: int = None):
    """Insert a new, empty code cell into the notebook and persist the
    change to disk.

    If after_index is given, the new cell is inserted immediately after
    that cell's position; otherwise it's appended at the end. Same index
    shifting problem as delete (everything from the insertion point on
    shifts up by one), so cached per-cell state is reset the same way.
    """
    nb = session["nb"]
    new_cell = nbformat.v4.new_code_cell(source="")
    if after_index is None:
        nb.cells.append(new_cell)
    else:
        insert_at = min(after_index + 1, len(nb.cells))
        nb.cells.insert(insert_at, new_cell)
    save_notebook(filepath, session)
    _reset_cached_cell_state(filepath)


# ============================= File previews ================================
def render_txt_preview(filepath: str):
    """Read-only preview of a plain text file."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
    except Exception as e:
        st.error(f"Could not read file: {e}")
        return
    st.caption(f"{len(text.splitlines())} lines")
    st.code(text, language="text")


def render_csv_preview(filepath: str):
    """Read-only tabular preview of a .csv file."""
    try:
        import pandas as pd
    except ImportError:
        st.error(
            "Previewing .csv files requires pandas. "
            "Install with: pip install pandas --break-system-packages"
        )
        return
    try:
        df = pd.read_csv(filepath)
    except Exception as e:
        st.error(f"Could not parse CSV: {e}")
        return
    st.caption(f"{len(df)} rows × {len(df.columns)} columns")
    st.dataframe(df, use_container_width=True)


def render_xlsx_preview(filepath: str):
    """Read-only tabular preview of an .xlsx file, with a sheet picker
    when the workbook has more than one sheet."""
    try:
        import pandas as pd
    except ImportError:
        st.error(
            "Previewing .xlsx files requires pandas + openpyxl. "
            "Install with: pip install pandas openpyxl --break-system-packages"
        )
        return
    try:
        sheets = pd.read_excel(filepath, sheet_name=None, engine="openpyxl")
    except Exception as e:
        st.error(f"Could not parse spreadsheet: {e}")
        return

    sheet_names = list(sheets.keys())
    if not sheet_names:
        st.info("This workbook has no sheets.")
        return

    if len(sheet_names) > 1:
        chosen = st.selectbox("Sheet", sheet_names, key=f"xlsx_sheet_{filepath}")
    else:
        chosen = sheet_names[0]

    df = sheets[chosen]
    st.caption(f"{len(df)} rows × {len(df.columns)} columns")
    st.dataframe(df, use_container_width=True)


def render_image_preview(filepath: str, ext: str):
    """Read-only preview for raster images (png/jpg/jpeg) and SVG."""
    if ext == ".svg":
        try:
            with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                svg_content = f.read()
        except Exception as e:
            st.error(f"Could not read SVG: {e}")
            return
        b64 = base64.b64encode(svg_content.encode("utf-8")).decode("ascii")
        st.markdown(
            f'<img src="data:image/svg+xml;base64,{b64}" '
            'style="max-width:100%; background:white; padding:8px; border-radius:4px;" />',
            unsafe_allow_html=True,
        )
    else:
        try:
            st.image(filepath, use_container_width=True)
        except Exception as e:
            st.error(f"Could not display image: {e}")


def render_pdf_preview(filepath: str):
    """Read-only preview of a PDF, embedded via a base64 data URL iframe,
    plus a download button."""
    try:
        with open(filepath, "rb") as f:
            data = f.read()
    except Exception as e:
        st.error(f"Could not read PDF: {e}")
        return

    st.caption(f"{len(data) / 1024:.1f} KB")

    b64 = base64.b64encode(data).decode("ascii")
    components.html(
        f'<iframe src="data:application/pdf;base64,{b64}" '
        'width="100%" height="800" style="border:none;"></iframe>',
        height=820,
    )

    st.download_button(
        "⬇️ Download PDF",
        data=data,
        file_name=os.path.basename(filepath),
        mime="application/pdf",
        use_container_width=True,
    )
def render_docx_preview(filepath: str):
    """Read-only preview of a Word document: paragraphs (with rough
    heading styling) followed by any tables."""
    try:
        import docx
    except ImportError:
        st.error(
            "Previewing .docx files requires python-docx. "
            "Install with: pip install python-docx --break-system-packages"
        )
        return
    try:
        document = docx.Document(filepath)
    except Exception as e:
        st.error(f"Could not read document: {e}")
        return

    any_content = False
    for para in document.paragraphs:
        text = para.text
        if not text.strip():
            continue
        any_content = True
        style = (para.style.name or "").lower() if para.style else ""
        if style == "title" or style.startswith("heading 1"):
            st.header(text)
        elif style.startswith("heading 2"):
            st.subheader(text)
        elif style.startswith("heading"):
            st.markdown(f"**{text}**")
        else:
            st.write(text)

    if document.tables:
        any_content = True
        st.caption(f"{len(document.tables)} table(s) in document")
        for i, table in enumerate(document.tables):
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            if rows:
                st.markdown(f"**Table {i + 1}**")
                st.table(rows)

    if not any_content:
        st.info("This document appears to be empty.")


def render_file_preview(filepath: str, ext: str):
    """Dispatch to the right read-only preview renderer for `ext`."""
    if ext == ".txt":
        render_txt_preview(filepath)
    elif ext == ".csv":
        render_csv_preview(filepath)
    elif ext == ".xlsx":
        render_xlsx_preview(filepath)
    elif ext in (".png", ".jpg", ".jpeg", ".svg"):
        render_image_preview(filepath, ext)
    elif ext == ".pdf":
        render_pdf_preview(filepath)
    elif ext == ".docx":
        render_docx_preview(filepath)
    else:
        st.info("No preview available for this file type.")


# ============================= Top navbar ====================================
def _status_badge(label: str, active: bool) -> str:
    """Small pill-shaped status indicator used in the navbar (e.g. process
    running, kernel running). Colour follows the same blue-accent /
    neutral-grey language already used for the selected-file highlight in
    the sidebar tree, so the whole UI reads as one consistent design."""
    color = "#2196f3" if active else "#8a8a8a"
    bg = "rgba(33,150,243,0.16)" if active else "rgba(128,128,128,0.14)"
    dot = "●" if active else "○"
    return (
        f'<span class="ide-badge" style="color:{color};background:{bg};'
        f'border-color:{color}33;">{dot}&nbsp;{label}</span>'
    )


def render_top_navbar(root_folder: str, selected_file: str | None):
    """Render a fixed-feeling, full-width app navbar: brand + version on
    the left, a breadcrumb of the current folder/file in the middle, and
    live status badges (running process / notebook kernel) on the right.
    Pure HTML/CSS (no components.html) so it costs nothing extra to
    re-render on every Streamlit rerun."""
    proc = st.session_state.proc
    proc_running = proc is not None and proc.poll() is None

    kernel_badge_html = ""
    if selected_file:
        ext = os.path.splitext(selected_file)[1].lower()
        if ext == ".ipynb":
            session = st.session_state.nb_sessions.get(selected_file)
            kernel_live = bool(session and session["kernel_started"])
            kernel_badge_html = _status_badge("Kernel", kernel_live)

    proc_badge_html = _status_badge("Process", proc_running)

    root_name = os.path.basename(root_folder.rstrip("/\\")) or root_folder
    if selected_file:
        try:
            file_rel = os.path.relpath(selected_file, root_folder)
        except ValueError:
            file_rel = os.path.basename(selected_file)
        breadcrumb = f"{root_name} <span class='ide-crumb-sep'>›</span> {file_rel}"
    else:
        breadcrumb = f"{root_name} <span class='ide-crumb-sep'>›</span> <em>no file selected</em>"

    st.markdown(
        f"""
        <style>
        .ide-navbar {{
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 16px;
            padding: 10px 18px;
            margin: -1rem -1rem 14px -1rem;
            background: linear-gradient(180deg, rgba(38,39,48,0.96) 0%, rgba(28,29,36,0.96) 100%);
            border-bottom: 1px solid rgba(255,255,255,0.08);
            box-shadow: 0 2px 8px rgba(0,0,0,0.25);
        }}
        .ide-navbar-brand {{
            display: flex;
            align-items: baseline;
            gap: 8px;
            white-space: nowrap;
        }}
        .ide-navbar-title {{
            font-size: 1.05em;
            font-weight: 700;
            color: #f0f0f0;
            letter-spacing: 0.01em;
        }}
        .ide-navbar-version {{
            font-size: 0.72em;
            font-weight: 600;
            color: #8a8a8a;
        }}
        .ide-navbar-crumb {{
            flex: 1;
            min-width: 0;
            text-align: center;
            font-size: 0.85em;
            font-family: "SFMono-Regular", ui-monospace, Menlo, Consolas, monospace;
            color: #c9c9c9;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
        }}
        .ide-crumb-sep {{
            color: #5a5a5a;
            margin: 0 4px;
        }}
        .ide-navbar-status {{
            display: flex;
            align-items: center;
            gap: 8px;
            white-space: nowrap;
        }}
        .ide-badge {{
            font-size: 0.72em;
            font-weight: 600;
            padding: 3px 9px;
            border-radius: 999px;
            border: 1px solid;
        }}
        </style>
        <div class="ide-navbar">
            <div class="ide-navbar-brand">
                <span style="font-size:1.2em;">{CONFIG.APP_ICON}</span>
                <span class="ide-navbar-title">{CONFIG.APP_NAME}</span>
                <span class="ide-navbar-version">v{CONFIG.VERSION}</span>
            </div>
            <div class="ide-navbar-crumb">{breadcrumb}</div>
            <div class="ide-navbar-status">{kernel_badge_html}{proc_badge_html}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


# ============================= Delete dialog ================================
@st.dialog("⚠️ Confirm delete")
def delete_item_dialog(root_folder: str, rel_path: str):
    if not rel_path.strip():
        st.warning("Enter a path to delete first.")
        if st.button("Close"):
            st.session_state.show_delete_dialog = False
            st.rerun()
        return

    target_name = os.path.basename(rel_path.rstrip("/\\")) or rel_path
    st.warning(
        f"This will permanently delete **{rel_path}** "
        "(and everything inside it, if it's a folder). This cannot be undone."
    )
    typed = st.text_input(f"Type **{target_name}** to confirm", key="delete_confirm_typed")

    col1, col2 = st.columns(2)
    with col1:
        confirm = st.button(
            "Delete permanently", type="primary", use_container_width=True,
            disabled=(typed.strip() != target_name),
        )
    with col2:
        cancel = st.button("Cancel", use_container_width=True)

    if cancel:
        st.session_state.show_delete_dialog = False
        st.session_state.pop("delete_confirm_typed", None)
        st.rerun()

    if confirm:
        ok, result = delete_path(root_folder, rel_path)
        if ok:
            _invalidate_listings()
            if st.session_state.selected_file == result:
                kill_current_process()
                st.session_state.selected_file = None
            st.session_state.show_delete_dialog = False
            st.session_state.pop("delete_confirm_typed", None)
            logger.info("Deleted path: %s", result)
            st.toast(f"Deleted: {rel_path}", icon="🗑️")
            st.rerun()
        else:
            logger.warning("Delete failed for %s: %s", rel_path, result)
            st.error(result)  # keep dialog open so the user can see what went wrong


# ============================= Sidebar: VS Code-style tree ==================
def _invalidate_listings():
    """Bump the filesystem-version counter, busting the st.cache_data
    cache backing the folder scan (see _scan_runnable_files/_scan_file_tree)
    so the next render re-walks the folder instead of showing stale data."""
    st.session_state.fs_version += 1


# A real non-breaking-space character (U+00A0), not the literal text
# "&nbsp;" — st.button labels are plain text, not HTML, so an HTML entity
# would just show up as the literal 6 characters "&nbsp;". Used here only
# as a small, reliable separator between a chevron/icon and the name that
# follows it (actual tree indentation is done with layout, see below).
NBSP = "\u00a0"

# Plain geometric triangles (U+25B8 / U+25BE), not emoji — these are basic
# glyphs present in every font, unlike "⌄"/"›" which can render inconsistently
# depending on the system font stack.
CHEVRON_COLLAPSED = "\u25b8"  # ▸
CHEVRON_EXPANDED = "\u25be"   # ▾

# Indentation is done with real layout (st.columns), not spaces/NBSP
# characters inside button labels. Text-based indentation is at the mercy
# of font metrics and never lines up precisely between different row types
# (a button vs. a plain div, an icon vs. no icon, etc). A narrow spacer
# column of a width proportional to depth gives every row — folder button,
# file button, or plain file row — an identical, pixel-accurate left edge
# for its content at a given depth.
TREE_LEVEL_WIDTH = 0.9   # relative width added per nesting level
TREE_BASE_GUTTER = 0.5   # small constant left gutter, even at depth 0
TREE_ROW_TOTAL = 20      # total relative width of a tree row
TREE_MAX_DEPTH = 8       # cap how far indent keeps growing on very deep trees


def _tree_columns(depth: int):
    """Return (spacer_column, content_column) sized so that every row at
    the same depth — regardless of whether it's a folder button, a file
    button, or a plain text row — starts its content at exactly the same
    horizontal position."""
    capped_depth = min(depth, TREE_MAX_DEPTH)
    indent_units = TREE_BASE_GUTTER + TREE_LEVEL_WIDTH * capped_depth
    content_units = max(TREE_ROW_TOTAL - indent_units, 6)
    return st.columns([indent_units, content_units], gap="small")


def inject_tree_alignment_script():
    """Force folder rows left-aligned and file rows centered via direct DOM
    manipulation, bypassing CSS entirely.

    Several rounds of CSS-only attempts (targeting the button element, a
    st.container(key=...) wrapper class, and a structural DOM-depth trick)
    all silently failed to move the text, because Streamlit wraps a
    button's label in its own inner container that can carry a competing
    text-align, and Streamlit's internal class names/markup shape have
    also shifted across versions — both defeat selector-guessing.

    This sidesteps all of that: it reads each button's own rendered text
    (which we fully control — folder rows start with a chevron, runnable
    and previewable file rows start with one of the known file-type
    icons) and sets inline `!important` styles directly on the button and
    everything inside it. Inline `!important` styles beat virtually any
    external stylesheet rule for that exact element. A MutationObserver
    re-applies this after every Streamlit rerun, since Streamlit replaces
    these DOM nodes each time.
    """
    file_icons = sorted(set(RUNNABLE_ICONS.values()) | set(PREVIEW_ICONS.values()))
    icons_js_array = "[" + ",".join(f'"{icon}"' for icon in file_icons) + "]"

    components.html(
        f"""
        <script>
        (function() {{
            var FILE_ICONS = {icons_js_array};
            function apply() {{
                var doc = window.parent.document;
                var sidebar = doc.querySelector('section[data-testid="stSidebar"]');
                if (!sidebar) return;
                var buttons = sidebar.querySelectorAll('button');
                buttons.forEach(function(btn) {{
                    var text = (btn.textContent || '').trim();
                    var isFolder = text.indexOf('{CHEVRON_COLLAPSED}') === 0 || text.indexOf('{CHEVRON_EXPANDED}') === 0;
                    var isFile = FILE_ICONS.some(function(icon) {{ return text.indexOf(icon) === 0; }});
                    if (!isFolder && !isFile) return;
                    var align = isFolder ? 'left' : 'center';
                    var justify = isFolder ? 'flex-start' : 'center';
                    var targets = [btn].concat(Array.prototype.slice.call(btn.querySelectorAll('div, p, span')));
                    targets.forEach(function(el) {{
                        el.style.setProperty('text-align', align, 'important');
                        el.style.setProperty('justify-content', justify, 'important');
                        el.style.setProperty('display', el === btn ? 'flex' : el.style.display || 'block', 'important');
                        el.style.setProperty('width', '100%', 'important');
                    }});
                }});
            }}
            apply();
            var root = window.parent.document.querySelector('section[data-testid="stSidebar"]') || window.parent.document.body;
            var timer = null;
            var observer = new MutationObserver(function() {{
                clearTimeout(timer);
                timer = setTimeout(apply, 50);
            }});
            observer.observe(root, {{childList: true, subtree: true}});
        }})();
        </script>
        """,
        height=0,
    )


def render_tree(subtree: dict, folder_root: str, rel_prefix: str, depth: int):
    """Recursively render a VS Code-style tree: folders first (with
    expand/collapse chevrons, no folder icon), then files, indented by
    depth. Clicking a .py/.ipynb file selects it in the editor/runner.
    Clicking a previewable file (pdf/docx/txt/xlsx/csv/png/jpg/jpeg/svg)
    selects it too, but opens a read-only preview instead. Any other file
    is shown for context as a plain, non-interactive muted row."""
    folder_names = sorted(subtree["__folders__"].keys(), key=str.lower)
    for name in folder_names:
        rel_path = f"{rel_prefix}/{name}" if rel_prefix else name
        is_expanded = rel_path in st.session_state.expanded_folders
        chevron = CHEVRON_EXPANDED if is_expanded else CHEVRON_COLLAPSED
        _, content_col = _tree_columns(depth)
        with content_col:
            # Full-width button -> hugs the left edge of the row (this is
            # also reinforced by the plain "justify-content: flex-start"
            # CSS rule below, which — unlike a st.container(key=...) based
            # rule — targets a real, static class Streamlit always emits,
            # so it isn't at the mercy of the container-key bug.
            if st.button(f"{chevron}{NBSP}{name}", key=f"tree_folder_{rel_path}", use_container_width=True):
                if is_expanded:
                    st.session_state.expanded_folders.discard(rel_path)
                else:
                    st.session_state.expanded_folders.add(rel_path)
                st.rerun()
        if is_expanded:
            render_tree(subtree["__folders__"][name], folder_root, rel_path, depth + 1)

    files = sorted(subtree["__files__"], key=lambda x: x[0].lower())
    for name, full_path in files:
        ext = os.path.splitext(name)[1].lower()
        is_runnable = ext in RUNNABLE_EXTS
        is_previewable = ext in PREVIEWABLE_EXTS
        is_selected = full_path == st.session_state.selected_file
        _, content_col = _tree_columns(depth)

        if is_runnable or is_previewable:
            icon = get_file_icon(name)
            btn_type = "primary" if is_selected else "secondary"
            with content_col:
                if st.button(
                    f"{icon}{NBSP}{name}", key=f"tree_file_{rel_prefix}_{name}",
                    use_container_width=True, type=btn_type,
                ):
                    if full_path != st.session_state.selected_file:
                        kill_current_process()
                    st.session_state.selected_file = full_path
                    st.rerun()
        else:
            # Plain HTML row (not a widget) since it's not interactive —
            # a small dot instead of a per-extension icon keeps a long
            # run of misc. files (data exports, configs, etc.) from
            # turning into a wall of mismatched glyphs.
            with content_col:
                st.markdown(
                    f"<div class='tree-inert-row'><span class='tree-dot'>•</span>{NBSP}{name}</div>",
                    unsafe_allow_html=True,
                )


with st.sidebar:
    st.markdown(
        """
        <style>
        /* Kill Streamlit's default vertical gap between stacked widgets so
           the tree rows sit flush against each other like a real explorer,
           instead of a loose stack of separated buttons. */
        section[data-testid="stSidebar"] div[data-testid="stVerticalBlock"] {
            gap: 0rem !important;
        }
        section[data-testid="stSidebar"] div.element-container {
            margin: 0 !important;
        }
        section[data-testid="stSidebar"] div[data-testid="stMarkdownContainer"] p {
            margin: 0 !important;
        }
        /* Each tree row is now a 2-column layout (spacer + content) for
           pixel-accurate indentation — make sure that row wrapper itself
           adds no extra vertical space between rows, and that its columns
           don't add unwanted top/bottom padding. */
        section[data-testid="stSidebar"] div[data-testid="stHorizontalBlock"] {
            margin: 0 !important;
            gap: 0.4rem !important;
            align-items: center !important;
        }
        section[data-testid="stSidebar"] div[data-testid="column"] {
            padding: 0 !important;
            margin: 0 !important;
        }

        /* Tree row buttons (folders + runnable files) — base/default is the
           FOLDER look: name left-aligned, chevron acting as an accordion
           toggle (▸ collapsed / ▾ expanded). File rows override this below
           to be centered instead. */
        section[data-testid="stSidebar"] div.stButton > button {
            width: 100% !important;
            justify-content: flex-start !important;
            text-align: left !important;
            padding: 3px 8px !important;
            margin: 0 !important;
            min-height: 26px !important;
            height: 26px !important;
            line-height: 20px !important;
            font-size: 0.85em !important;
            font-family: inherit !important;
            border: none !important;
            border-left: 3px solid transparent !important;
            border-radius: 0 !important;
            background: transparent !important;
            color: inherit !important;
            white-space: nowrap !important;
            overflow: hidden !important;
            text-overflow: ellipsis !important;
            display: block !important;
            box-shadow: none !important;
            transition: background 0.12s ease, border-color 0.12s ease !important;
        }
        section[data-testid="stSidebar"] div.stButton > button:hover {
            background: rgba(128,128,128,0.18) !important;
            color: inherit !important;
        }
        section[data-testid="stSidebar"] div.stButton > button:focus:not(:active) {
            box-shadow: none !important;
        }
        section[data-testid="stSidebar"] div.stButton > button[kind="primary"] {
            background: rgba(33,150,243,0.22) !important;
            border-left: 3px solid #2196f3 !important;
            font-weight: 600 !important;
        }
        section[data-testid="stSidebar"] div.stButton > button[kind="primary"]:hover {
            background: rgba(33,150,243,0.32) !important;
        }

        /* NOTE: text-align / justify-content for folder vs. file rows is
           NOT handled here — see inject_tree_alignment_script() below.
           Streamlit wraps a button's label in its own inner container
           (typically div[data-testid="stMarkdownContainer"] > p), which
           can carry its own text-align. An ancestor rule on <button> loses
           to an explicit rule on that descendant regardless of !important
           (inheritance vs. cascade are separate), and Streamlit's internal
           class names/DOM shape have also changed across versions — both
           of which made every CSS-only attempt at this silently no-op.
           A small JS pass that reads each button's actual rendered text
           and force-sets inline !important styles sidesteps all of that. */

        /* Non-runnable files: plain dimmed row, same box model as a button
           row (including the 3px transparent left border) so it lines up
           exactly with folders/runnable files above and below it. Centered
           to match the runnable-file rows. */
        .tree-inert-row {
            padding: 3px 8px 3px 11px;
            min-height: 26px;
            box-sizing: border-box;
            line-height: 20px;
            font-size: 0.85em;
            color: #8a8a8a;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
            text-align: center;
        }
        .tree-dot {
            color: #9a9a9a;
            font-size: 0.7em;
            vertical-align: middle;
        }

        /* Root folder name above the toolbar row — single line, truncated,
           small-caps style like a VS Code panel header. */
        .tree-root-title {
            font-size: 0.78em;
            font-weight: 700;
            letter-spacing: 0.03em;
            color: #b0b0b0;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
            margin: 20px;
        }

        /* Toolbar icon buttons (new file / new folder / refresh / collapse) */
        section[data-testid="stSidebar"] div[data-testid="column"] div.stButton > button {
            text-align: center !important;
            justify-content: center !important;
            padding: 2px 0 !important;
            min-height: 28px !important;
            height: 28px !important;
            white-space: nowrap !important;
            border-left: none !important;
            border-radius: 4px !important;
        }
        .e1wbovuq0{
        
            margin-bottom: 0px;
            margin-top: 10px;
            width: 100%;
            border-style: solid;
            border-width: 1px;
            border-color: rgba(250, 250, 250, 0.2);
            border-radius: 0.5rem;
            }
        
        </style>
        """,
        unsafe_allow_html=True,
    )

    # ---- Folder selection: read-only path display + native "Pick folder" ----
    path_cols = st.columns([5, 1])
    with path_cols[0]:
        st.text_input(
            "Folder path", value=st.session_state.folder_path,
            label_visibility="collapsed", disabled=True,
        )
    with path_cols[1]:
        pick_clicked = st.button("📂", key="pick_folder", help="Pick a folder", use_container_width=True)

    if pick_clicked:
        ok, result = browse_for_folder(st.session_state.folder_path)
        if ok:
            st.session_state.folder_path = result
            st.session_state.selected_file = None
            st.session_state.expanded_folders = set()
            _invalidate_listings()
            st.toast(f"Folder selected: {result}", icon="📂")
            st.rerun()
        elif result:  # non-empty result on failure means a real error (empty = user just cancelled)
            st.error(result)

    folder = st.session_state.folder_path

    st.divider()

    # ---- VS Code-style toolbar: root name on its own line (truncated, never
    # wraps), then a separate row of small evenly-spaced icon buttons below —
    # splitting these into two rows is what keeps a long folder name and four
    # icon buttons from ever fighting each other for space and wrapping.
    root_name = os.path.basename(folder.rstrip("/\\")) or folder
    # st.markdown(f"<div class='tree-root-title'>{root_name.upper()}</div>", unsafe_allow_html=True)

    hcol1, hcol2, hcol3, hcol4 = st.columns(4)
    with hcol1:
        if st.button("📄+", key="toolbar_new_file", help="New file", use_container_width=True):
            st.session_state.show_inline_new_file = not st.session_state.show_inline_new_file
            st.session_state.show_inline_new_folder = False
    with hcol2:
        if st.button("📁+", key="toolbar_new_folder", help="New folder", use_container_width=True):
            st.session_state.show_inline_new_folder = not st.session_state.show_inline_new_folder
            st.session_state.show_inline_new_file = False
    with hcol3:
        if st.button("🔄", key="toolbar_refresh", help="Refresh", use_container_width=True):
            _invalidate_listings()
            st.rerun()
    with hcol4:
        if st.button("⊟", key="toolbar_collapse", help="Collapse all", use_container_width=True):
            st.session_state.expanded_folders = set()
            st.rerun()

    # ---- Inline "path input box" for creating a new file, VS Code style ----
    if st.session_state.show_inline_new_file:
        with st.form(key="inline_new_file_form", clear_on_submit=False):
            new_file_path = st.text_input(
                "New file path (relative to root)",
                placeholder="e.g. scripts/new_script.py",
                key="inline_new_file_input",
            )
            fcols = st.columns(2)
            with fcols[0]:
                create_file_submit = st.form_submit_button("Create", type="primary", use_container_width=True)
            with fcols[1]:
                cancel_file_submit = st.form_submit_button("Cancel", use_container_width=True)
        if cancel_file_submit:
            st.session_state.show_inline_new_file = False
            st.rerun()
        if create_file_submit:
            ok, result = create_new_file(folder, new_file_path)
            if ok:
                _invalidate_listings()
                st.session_state.selected_file = result
                st.session_state.show_inline_new_file = False
                st.toast(f"Created: {os.path.relpath(result, folder)}", icon="✅")
                st.rerun()
            else:
                st.error(result)

    # ---- Inline "path input box" for creating a new folder ----
    if st.session_state.show_inline_new_folder:
        with st.form(key="inline_new_folder_form", clear_on_submit=False):
            new_folder_path = st.text_input(
                "New folder path (relative to root)",
                placeholder="e.g. data/raw",
                key="inline_new_folder_input",
            )
            
            fcols = st.columns(2)
            with fcols[0]:
                create_folder_submit = st.form_submit_button("Create", type="primary", use_container_width=True)
            with fcols[1]:
                cancel_folder_submit = st.form_submit_button("Cancel", use_container_width=True)
        if cancel_folder_submit:
            st.session_state.show_inline_new_folder = False
            st.rerun()
        if create_folder_submit:
            ok, result = create_new_folder(folder, new_folder_path)
            if ok:
                _invalidate_listings()
                st.session_state.show_inline_new_folder = False
                st.toast(f"Created: {os.path.relpath(result, folder)}", icon="✅")
                st.rerun()
            else:
                st.error(result)
        
    # with st.expander("🗑️ Delete a file / folder"):
    #     default_delete_path = ""
    #     if st.session_state.selected_file and os.path.isdir(folder):
    #         try:
    #             default_delete_path = os.path.relpath(st.session_state.selected_file, folder)
    #         except ValueError:
    #             default_delete_path = ""
    #     delete_rel_path_input = st.text_input(
    #         "Path relative to root folder",
    #         value=default_delete_path,
    #         key="delete_rel_path_input",
    #         placeholder="e.g. old_script.py  or  old_folder",
    #     )
    #     if st.button("Delete…", use_container_width=True):
    #         st.session_state.delete_target_rel_path = delete_rel_path_input
    #         st.session_state.show_delete_dialog = True

    # if st.session_state.get("show_delete_dialog"):
    #     delete_item_dialog(folder, st.session_state.get("delete_target_rel_path", ""))

    st.divider()

    if os.path.isdir(folder):
        fs_version = st.session_state.fs_version
        files = list_runnable_files(folder, fs_version)
        tree = build_file_tree(folder, fs_version)

        # Root starts expanded so the top-level contents are visible
        # immediately, like a freshly opened VS Code workspace. Track the
        # last folder we did this for so switching folders re-expands root
        # without re-collapsing it on every unrelated rerun.
        if st.session_state.get("_last_tree_folder") != folder:
            st.session_state.expanded_folders.add("")
            st.session_state["_last_tree_folder"] = folder

        if not files:
            st.info("No files found in this folder.")
        else:
            st.caption(f"{len(files)} runnable files")
            render_tree(tree, folder, "", depth=0)
            inject_tree_alignment_script()
    else:
        st.error("That folder path doesn't exist.")

    st.divider()
    with st.expander("🪵 Application log", expanded=False):
        log_path = CONFIG.LOG_DIR / CONFIG.LOG_FILE
        if log_path.exists():
            try:
                tail_lines = log_path.read_text(encoding="utf-8", errors="replace").splitlines()[-40:]
                st.code("\n".join(tail_lines) or "(log is empty)", language="text")
            except OSError as e:
                st.caption(f"Could not read log file: {e}")
        else:
            st.caption("No log entries yet.")
        st.caption(f"Log file: `{log_path}`")

# st.divider()

selected = st.session_state.selected_file
render_top_navbar(st.session_state.folder_path, selected)

if not selected:
    st.info("👈 Enter a folder path in the sidebar, scan it, then pick a file from the tree.")
    st.stop()
selected_ext = os.path.splitext(selected)[1].lower()
is_py = selected_ext == ".py"
is_ipynb = selected_ext == ".ipynb"
is_preview_only = selected_ext in PREVIEWABLE_EXTS

try:
    if is_py:
        # ============================= Layout: buttons -> code -> console ========
        # (stacked vertically, Jupyter-style, instead of side-by-side columns)
        st.markdown("**📝 Code** (edit, then Save and/or Run)")

        btn_cols = st.columns(4)
        with btn_cols[0]:
            save_clicked = st.button("💾 Save", use_container_width=True)
        with btn_cols[1]:
            run_clicked = st.button("▶ Save & Run", type="primary", use_container_width=True)
        with btn_cols[2]:
            revert_clicked = st.button("↩ Revert", use_container_width=True)
        with btn_cols[3]:
            delete_file_clicked = st.button("🗑️ Delete file", use_container_width=True)

        content = load_file_content(selected)
        editor_height = ace_height_for(content)
        # Bumped on Revert to force the editor component to fully remount with
        # the reverted text — custom components like streamlit-ace keep their
        # own internal buffer and otherwise ignore a changed `value` prop on
        # rerun, which is why Revert previously appeared to do nothing.
        editor_ver = st.session_state.editor_version.get(selected, 0)

        if ACE_AVAILABLE:
            new_content = st_ace(
                value=content,
                language="python",
                theme="tomorrow_night",
                keybinding="vscode",
                key=f"ace_{selected}_v{editor_ver}",
                height=editor_height,
                font_size=13,
                show_gutter=True,
                wrap=False,
                # auto_update=True keeps the buffer synced on every keystroke.
                # auto_update=False relies on Ctrl+Enter or blur to commit,
                # which is unreliable across browsers/OSes (e.g. Cmd+Enter on
                # macOS doesn't always fire) — auto_update=True sidesteps that
                # entirely so Save/Run/Revert always see the latest text.
                auto_update=True,
            )
        else:
            new_content = st.text_area(
                "Code", value=content, height=editor_height,
                key=f"ta_{selected}_v{editor_ver}", label_visibility="collapsed",
            )
        st.session_state.editor_content[selected] = new_content

        if save_clicked:
            save_file_content(selected, new_content)
            st.toast("Saved.", icon="💾")

        if revert_clicked:
            original = st.session_state.original_content.get(selected, "")
            st.session_state.editor_content[selected] = original
            st.session_state.editor_version[selected] = editor_ver + 1
            st.rerun()

        if run_clicked:
            save_file_content(selected, new_content)
            start_process(selected)
            st.rerun()

        if delete_file_clicked:
            # Reuse the same type-to-confirm dialog as the sidebar's delete
            # flow. It's rendered from within the sidebar block, so we just
            # set the target + flag here and rerun to let it pick this up.
            st.session_state.delete_target_rel_path = os.path.relpath(
                selected, st.session_state.folder_path
            )
            st.session_state.show_delete_dialog = True
            st.rerun()

        st.divider()
        st.markdown("**▶ Console / Output**")

        def render_console():
            if "proc" not in st.session_state:
                st.session_state.proc = None
                st.session_state.proc_output = ""
                st.session_state.proc_returncode = None
                st.session_state.proc_file = None
            poll_process_output()
            output = st.session_state.proc_output
            st.code(output if output else "(no output yet — click Save & Run)", language="text")

            proc = st.session_state.proc
            still_running = proc is not None and st.session_state.proc_returncode is None

            if still_running:
                st.info("⏳ Running... if the script is waiting for input(), type below and hit Send.")
                with st.form(key=f"stdin_form_{selected}", clear_on_submit=True):
                    fcols = st.columns([5, 1])
                    with fcols[0]:
                        user_input = st.text_input(
                            "stdin", label_visibility="collapsed",
                            placeholder="Type input for the running script...",
                        )
                    with fcols[1]:
                        submitted = st.form_submit_button("Send")
                if submitted and user_input != "":
                    send_stdin(user_input)
                    st.rerun()

                if st.button("⏹ Stop process", key=f"stop_{selected}"):
                    kill_current_process()
                    st.rerun()
            else:
                rc = st.session_state.proc_returncode
                if rc is None:
                    pass
                elif rc == 0:
                    st.success("✅ Finished successfully (exit code 0).")
                else:
                    st.error(f"❌ Process exited with code {rc}. See traceback above.")

        # Only auto-poll on a timer while something is actually running. Doing
        # this unconditionally (even when idle / no process) burns CPU and
        # websocket bandwidth forever in the background, which is the biggest
        # source of sluggishness in the whole app.
        _proc = st.session_state.proc
        _is_running = _proc is not None and _proc.poll() is None
        if _is_running:
            st.fragment(run_every=CONFIG.CONSOLE_POLL_SECONDS)(render_console)()
        else:
            render_console()

    elif is_ipynb:
        if not NBCLIENT_AVAILABLE:
            st.error("nbclient / nbformat not installed. Run: pip install nbclient nbformat --break-system-packages")
            st.stop()

        top_cols = st.columns([1, 1, 4])
        with top_cols[0]:
            if st.button("🔄 Restart Kernel", use_container_width=True):
                restart_kernel(selected)
                st.toast("Kernel restarted — shared state (variables) has been reset.", icon="🔄")
                st.rerun()
        with top_cols[1]:
            session_exists = selected in st.session_state.nb_sessions
            kernel_live = session_exists and st.session_state.nb_sessions[selected]["kernel_started"]
            # st.caption(f"Kernel: {'🟢 running' if kernel_live else '⚪ not started'}")
            kernel_status_label = "Running" if kernel_live else "Not started"
            kernel_status_color = "#2196f3" if kernel_live else "#f4f0f0"
            kernel_status_bg = "rgba(33,150,243,0.15)" if kernel_live else "rgba(128,128,128,0.12)"

            st.markdown(
                f"""
                <div style="
                    display: inline-flex;
                    align-items: center;
                    justify-content: center;
                    width: 100%;
                    min-height: 38px;
                    height: auto;
                    padding: 2px 8px;
                    font-size: 0.85em;
                    font-weight: 600;
                    border-radius: 4px;
                    background: {kernel_status_bg};
                    border-left: 3px solid {kernel_status_color};
                    color: inherit;
                    box-sizing: border-box;
                ">
                    Kernel: {kernel_status_label}
                </div>
                """,
                unsafe_allow_html=True,
            )

        session = get_or_create_nb_session(selected)
        nb = session["nb"]

        code_cell_indices = [i for i, c in enumerate(nb.cells) if c.cell_type == "code"]

        # ---- "Add cell" when the notebook has no code cells at all ----
        if not code_cell_indices:
            st.info("This notebook has no code cells yet.")
            if st.button("➕ Add first cell", use_container_width=True):
                add_notebook_cell(selected, session, after_index=None)
                st.toast("Cell added.", icon="➕")
                st.rerun()

        for pos, idx in enumerate(code_cell_indices):
            cell = nb.cells[idx]
            src_key = (selected, idx)
            current_src = st.session_state.nb_cell_source.get(src_key, cell.source)
            comment_only = is_comment_only(current_src)

            cell_label = f"**Cell [{pos + 1}/{len(code_cell_indices)}]**"
            if comment_only:
                cell_label += "  ·  _comment only_"
            st.markdown(cell_label)

            # ==================== Stacked, Jupyter-style: ====================
            # action buttons -> code editor -> output/error, all in one column
            cell_save = cell_run = cell_clear = cell_revert = False
            if comment_only:
                # Nothing to run or save in a comment-only cell — the buttons
                # left are add (a new cell below) and delete.
                b_add, b_delete = st.columns(2)
                with b_add:
                    cell_add = st.button(
                        "➕ Add cell below", key=f"add_{selected}_{idx}", use_container_width=True,
                    )
                with b_delete:
                    cell_delete = st.button(
                        "🗑️ Delete cell", key=f"delete_{selected}_{idx}", use_container_width=True,
                    )
            else:
                b1, b2, b3, b4, b5, b6 = st.columns(6)
                with b1:
                    cell_save = st.button("💾 Save", key=f"save_{selected}_{idx}", use_container_width=True)
                with b2:
                    cell_run = st.button("▶ Run", key=f"run_{selected}_{idx}", use_container_width=True, type="primary")
                with b3:
                    cell_clear = st.button("🗑 Clear Output", key=f"clear_{selected}_{idx}", use_container_width=True)
                with b4:
                    cell_revert = st.button("↩ Revert", key=f"revert_{selected}_{idx}", use_container_width=True)
                with b5:
                    cell_add = st.button("➕ Add cell below", key=f"add_{selected}_{idx}", use_container_width=True)
                with b6:
                    cell_delete = st.button("🗑️ Delete cell", key=f"delete_{selected}_{idx}", use_container_width=True)

            if cell_add:
                add_notebook_cell(selected, session, after_index=idx)
                st.toast(f"Cell added after {pos + 1}.", icon="➕")
                st.rerun()

            if cell_delete:
                delete_notebook_cell(selected, session, idx)
                st.toast(f"Deleted cell {pos + 1}.", icon="🗑️")
                st.rerun()

            cell_editor_height = ace_height_for(current_src, min_lines=3)
            # Bumped on Revert to force this cell's editor to remount with the
            # reverted text (same remount trick as the .py file editor above).
            cell_ver = st.session_state.cell_editor_version.get(src_key, 0)

            if ACE_AVAILABLE:
                edited = st_ace(
                    value=current_src,
                    language="python",
                    theme="tomorrow_night",
                    keybinding="vscode",
                    key=f"ace_{selected}_{idx}_v{cell_ver}",
                    height=cell_editor_height,
                    font_size=13,
                    show_gutter=True,
                    wrap=False,
                    # See the .py editor above: auto_update=True keeps the
                    # buffer synced on every keystroke instead of depending on
                    # a Ctrl/Cmd+Enter commit shortcut that isn't reliable
                    # across platforms.
                    auto_update=True,
                )
            else:
                edited = st.text_area(
                    f"cell_{idx}", value=current_src, height=cell_editor_height,
                    key=f"ta_{selected}_{idx}_v{cell_ver}", label_visibility="collapsed",
                )
            st.session_state.nb_cell_source[src_key] = edited

            out_placeholder = st.container()

            if cell_save:
                cell.source = edited
                save_notebook(selected, session)
                st.session_state.nb_cell_original[src_key] = edited
                st.toast(f"Cell {pos + 1} saved.", icon="💾")

            if cell_clear:
                cell.outputs = []
                cell.execution_count = None
                st.rerun()

            if cell_revert:
                original = st.session_state.nb_cell_original.get(src_key, cell.source)
                st.session_state.nb_cell_source[src_key] = original
                st.session_state.cell_editor_version[src_key] = cell_ver + 1
                st.rerun()

            if cell_run:
                cell.source = edited
                save_notebook(selected, session)  # persist the code that's about to run
                st.session_state.nb_cell_original[src_key] = edited
                with st.spinner(f"Running cell {pos + 1}..."):
                    ok = run_single_cell(session, idx, edited)
                save_notebook(selected, session)  # persist the resulting outputs too
                if ok:
                    st.toast(f"Cell {pos + 1} ran successfully.", icon="✅")
                else:
                    st.toast(f"Cell {pos + 1} raised an error.", icon="❌")

            if cell.outputs:
                for output in cell.outputs:
                    render_notebook_output(output, out_placeholder)
            else:
                out_placeholder.caption("(no output yet)")

            st.divider()

        # ---- "Add cell" at the very end, so you can append without having to
        # use a specific existing cell's "Add cell below" button ----
        if code_cell_indices:
            if st.button("➕ Add cell at end", use_container_width=True):
                add_notebook_cell(selected, session, after_index=None)
                st.toast("Cell added at end.", icon="➕")
                st.rerun()

    elif is_preview_only:
        # Read-only preview for pdf / docx / txt / xlsx / csv / png / jpg /
        # jpeg / svg files clicked in the tree. No editing or running here —
        # just a quick look at the content, plus a delete option for parity
        # with the runnable-file view.
        top_cols = st.columns([3, 1])
        with top_cols[0]:
            st.markdown(f"**{PREVIEW_ICONS.get(selected_ext, '📄')} Preview** (read-only)")
        with top_cols[1]:
            delete_preview_clicked = st.button("🗑️ Delete file", use_container_width=True)

        if delete_preview_clicked:
            st.session_state.delete_target_rel_path = os.path.relpath(
                selected, st.session_state.folder_path
            )
            st.session_state.show_delete_dialog = True
            st.rerun()

        st.divider()
        render_file_preview(selected, selected_ext)

    else:
        st.info("No preview available for this file type. Click a .py, .ipynb, or previewable file in the sidebar.")
except Exception as e:
    logger.exception("Unhandled error while rendering the main content area for %s", selected)
    st.error("Something went wrong while rendering this file. The details have been logged.")
    with st.expander("Show error details"):
        st.exception(e)